import os
import docx
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException
from app.models import Dataset, DocxDocument, DocxDocumentCreate
# 如果还需要响应 schema：
# from app.models import DocxDocumentResponse
from app.database import get_db
from app.models import Dataset, DocxDocument, DocxDocumentCreate

class DocxService:
    @staticmethod
    async def parse_docx(file_path: str) -> str:
        """解析docx文件内容"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"无法解析DOCX文件: {str(e)}")
    
    @staticmethod
    async def upload_docx_to_db(file_path: str, filename: str, dataset_id: Optional[int] = None) -> Dict[str, Any]:
        """上传单个docx文件到数据库"""
        try:
            # 解析docx文件
            content = await DocxService.parse_docx(file_path)
            
            # 存储到数据库
            db = next(get_db())
            docx_doc = DocxDocument(
                filename=filename,
                content=content,
                dataset_id=dataset_id
            )
            db.add(docx_doc)
            db.commit()
            db.refresh(docx_doc)
            
            return {
                "id": docx_doc.id,
                "filename": docx_doc.filename,
                "content_length": len(content),
                "dataset_id": dataset_id
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"上传DOCX文件失败: {str(e)}")
    
    @staticmethod
    async def batch_upload_docx_files(dataset_name: str, dataset_description: Optional[str], files_info: List[Dict[str, str]]) -> Dict[str, Any]:
        """批量上传docx文件"""
        db = next(get_db())
        
        # 创建或获取数据集
        dataset = db.query(Dataset).filter(Dataset.name == dataset_name).first()
        if not dataset:
            dataset = Dataset(name=dataset_name, description=dataset_description)
            db.add(dataset)
            db.commit()
            db.refresh(dataset)
        
        # 批量上传文件
        upload_results = []
        for file_info in files_info:
            file_path = file_info.get("file_path")
            
            # 验证文件路径
            if not os.path.exists(file_path):
                upload_results.append({
                    "file_path": file_path,
                    "status": "error",
                    "message": "文件不存在"
                })
                continue
                
            # 获取文件名
            filename = os.path.basename(file_path)
            
            try:
                # 上传文件到数据库
                result = await DocxService.upload_docx_to_db(file_path, filename, dataset.id)
                upload_results.append({
                    "file_path": file_path,
                    "status": "success",
                    "document_id": result["id"],
                    "filename": result["filename"]
                })
            except Exception as e:
                upload_results.append({
                    "file_path": file_path,
                    "status": "error",
                    "message": str(e)
                })
        
        return {
            "dataset": {
                "id": dataset.id,
                "name": dataset.name,
                "description": dataset.description
            },
            "upload_results": upload_results
        }
    
    @staticmethod
    def get_docx_documents(
        dataset_id: Optional[int] = None, 
        skip: int = 0, 
        limit: int = 100, 
        filters: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """获取docx文档列表，支持分页和过滤"""
        db = next(get_db())
        query = db.query(DocxDocument)
        
        # 应用数据集过滤
        if dataset_id:
            query = query.filter(DocxDocument.dataset_id == dataset_id)
        
        # 应用其他过滤条件
        if filters:
            for key, value in filters.items():
                if hasattr(DocxDocument, key):
                    query = query.filter(getattr(DocxDocument, key) == value)
        
        # 获取总数
        total = query.count()
        
        # 应用分页
        documents = query.offset(skip).limit(limit).all()
        
        return {
            "total": total,
            "documents": [
                {
                    "id": doc.id,
                    "filename": doc.filename,
                    "content": doc.content,
                    "created_at": doc.created_at,
                    "dataset_id": doc.dataset_id
                }
                for doc in documents
            ]
        }
    
    @staticmethod
    def get_docx_document(document_id: int) -> Dict[str, Any]:
        """获取单个docx文档"""
        db = next(get_db())
        document = db.query(DocxDocument).filter(DocxDocument.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="文档不存在")
        
        return {
            "id": document.id,
            "filename": document.filename,
            "content": document.content,
            "created_at": document.created_at,
            "dataset_id": document.dataset_id
        }

docx_service = DocxService()